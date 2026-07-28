"""Executive workflows must do the whole chain — and stop before sending."""

from datetime import date, timedelta

import pytest
from docx import Document

from emefa.domain.crm import CrmRepository
from emefa.domain.documents import DocumentStore
from emefa.domain.profiles import ProfileRepository
from emefa.domain.tasks import TaskRepository
from emefa.domain.workflows import PROPOSAL_FOLLOW_UP_DAYS, WorkflowEngine

TODAY = date(2026, 7, 28)


def engine(tmp_path):
    database = tmp_path / "workflows.db"
    profiles = ProfileRepository(database)
    profiles.update_business(
        {"company_name": "Horizon SARL", "owner_name": "Koffi Gava", "preferred_name": "M. Gava"}
    )
    crm = CrmRepository(database)
    documents = DocumentStore(database)
    tasks = TaskRepository(database)
    return profiles, crm, documents, tasks, WorkflowEngine(profiles, crm, documents, tasks)


def test_commercial_proposal_runs_the_whole_chain(tmp_path):
    _, crm, documents, tasks, workflows = engine(tmp_path)
    client = crm.save_contact(
        name="Ama Mensah", kind="client", company="Mensah Logistics", email="ama@mensah.tg"
    )
    crm.save_deal(title="Mission précédente", contact_id=client.contact_id, stage="accepté")
    crm.log_interaction("Appel de cadrage", contact_id=client.contact_id, occurred_at="2026-07-20")

    result = workflows.commercial_proposal(
        client="Ama Mensah",
        subject="Refonte du site web",
        context="Le site actuel ne convertit plus.",
        items=[
            {"label": "Conception", "quantity": 1, "unit_price": 800_000},
            {"label": "Développement", "quantity": 2, "unit_price": 350_000},
        ],
        today=TODAY,
    )

    assert result["status"] == "prepared"
    assert result["total"] == 1_500_000
    steps = {step["name"]: step["status"] for step in result["steps"]}
    assert steps == {
        "client": "done", "historique": "done", "document": "done",
        "devis": "done", "relance": "done", "email": "done",
    }

    # The client was recognised, not duplicated.
    assert result["client"]["contact_id"] == client.contact_id
    history_step = next(s for s in result["steps"] if s["name"] == "historique")
    assert len(history_step["data"]["previous_deals"]) == 1
    assert len(history_step["data"]["last_interactions"]) >= 1

    # A real, priced, editable document.
    text = "\n".join(
        p.text for p in Document(str(documents.get(result["document"]["document_id"]))).paragraphs
    )
    assert "Proposition commerciale — Refonte du site web" in text
    assert "Mensah Logistics" in text
    tables = Document(str(documents.get(result["document"]["document_id"]))).tables
    assert tables[0].rows[0].cells[0].text == "Prestation"

    # The quotation is tracked so it can be chased later.
    deal = crm.get_deal(result["deal_id"])
    assert deal.amount == 1_500_000
    assert deal.stage == "brouillon"
    assert deal.document_id == result["document"]["document_id"]

    # The follow-up the executive would have forgotten.
    follow_up = tasks.get(result["follow_up_task_id"])
    assert follow_up.due_date == (TODAY + timedelta(days=PROPOSAL_FOLLOW_UP_DAYS)).isoformat()

    # Prepared, never sent.
    action = result["proposed_action"]
    assert action["tool"] == "email_send"
    assert action["requires_approval"] is True
    assert action["arguments"]["to"] == "ama@mensah.tg"
    assert "1 500 000" in action["arguments"]["body"]
    assert "Rien n'a été envoyé" in result["note"]


def test_unknown_client_is_created_and_missing_email_is_reported(tmp_path):
    _, crm, _, _, workflows = engine(tmp_path)
    result = workflows.commercial_proposal(
        client="Nouvelle Enseigne", subject="Audit", amount=400_000, today=TODAY
    )
    client_step = next(step for step in result["steps"] if step["name"] == "client")
    assert client_step["data"]["created"] is True
    assert crm.get_contact(result["client"]["contact_id"]).kind == "prospect"

    email_step = next(step for step in result["steps"] if step["name"] == "email")
    assert email_step["status"] == "skipped"
    # No address means no send is proposed — the gap is named instead.
    assert result["proposed_action"]["tool"] != "email_send"
    assert "e-mail" in result["proposed_action"]["label"]


def test_follow_up_workflow_uses_the_relationship_history(tmp_path):
    _, crm, _, tasks, workflows = engine(tmp_path)
    client = crm.save_contact(name="Horizon Group", kind="client", email="dg@horizon.tg")
    crm.save_deal(
        title="Phase 2", contact_id=client.contact_id, stage="envoyé", sent_at="2026-07-01"
    )
    crm.log_interaction("Réunion de cadrage", contact_id=client.contact_id, occurred_at="2026-06-15")

    result = workflows.follow_up(client="horizon", today=TODAY)
    assert result["status"] == "prepared"
    assert [deal["title"] for deal in result["pending_deals"]] == ["Phase 2"]
    assert result["history"]
    assert tasks.get(result["task_id"]) is not None
    body = result["proposed_action"]["arguments"]["body"]
    assert "Phase 2" in body
    assert "M. Gava" in body
    assert result["proposed_action"]["requires_approval"] is True


def test_follow_up_reports_an_unknown_client_instead_of_inventing_one(tmp_path):
    _, _, _, _, workflows = engine(tmp_path)
    result = workflows.follow_up(client="Personne", today=TODAY)
    assert result == {"workflow": "relance", "status": "failed", "error": "client_introuvable"}


@pytest.mark.asyncio
async def test_proposal_skill_stops_at_the_approval_gate(tmp_path):
    """The workflow prepares; sending still needs the user's explicit yes."""
    from emefa.config import Settings
    from emefa.domain.agent import AgentStep, RequestedAction
    from emefa.main import create_app

    class ScriptedBrain:
        def __init__(self):
            self.calls = 0

        async def think(self, history, tools):
            self.calls += 1
            if self.calls == 1:
                return AgentStep(
                    action=RequestedAction(
                        name="workflow_commercial_proposal",
                        arguments={"client": "Ama Mensah", "subject": "Audit", "amount": 500000},
                    )
                )
            return AgentStep(answer="La proposition est prête, je peux l'envoyer si vous validez.")

    app = create_app(Settings(database_path=tmp_path / "api.db"), brain=ScriptedBrain())
    app.state.crm.save_contact(name="Ama Mensah", kind="client", email="ama@mensah.tg")
    token = app.state.devices.enroll("Claude")[1]

    import httpx

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        run = await client.post(
            "/v1/agent/runs",
            json={"message": "Prépare une proposition commerciale pour Ama Mensah"},
            headers={"Authorization": f"Bearer {token}"},
        )

    # The whole chain ran without ever reaching a COMMUNICATE action.
    assert run.json()["status"] == "completed"
    assert len(app.state.crm.list_deals()) == 1
    assert len(app.state.tasks.list_open()) == 1
    assert len(app.state.documents.list()) == 1
