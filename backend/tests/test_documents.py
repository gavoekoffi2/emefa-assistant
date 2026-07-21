import io

import httpx
import pytest
from docx import Document

from emefa.config import Settings
from emefa.domain.agent import AgentStep, RequestedAction
from emefa.main import create_app


class ScriptedBrain:
    def __init__(self, action: RequestedAction):
        self.action = action
        self.calls = 0

    async def think(self, history, tools):
        self.calls += 1
        if self.calls == 1:
            return AgentStep(action=self.action)
        result = next(item for item in reversed(history) if item.get("role") == "tool")
        return AgentStep(answer=f"Document prêt : {result['content']['download_url']}")


async def _client_for(app):
    transport = httpx.ASGITransport(app=app)
    return httpx.AsyncClient(transport=transport, base_url="http://test")


@pytest.mark.asyncio
async def test_document_create_produces_downloadable_real_docx(tmp_path):
    brain = ScriptedBrain(RequestedAction(
        name="document_create",
        arguments={"title": "Compte rendu", "content": "Décision principale\nProchaine étape"},
    ))
    app = create_app(Settings(database_path=tmp_path / "emefa.db"), brain=brain)
    token = app.state.devices.enroll("Claude")[1]
    headers = {"Authorization": f"Bearer {token}"}

    async with await _client_for(app) as client:
        run = await client.post("/v1/agent/runs", json={"message": "Crée le compte rendu"}, headers=headers)
        assert run.status_code == 200
        assert run.json()["status"] == "completed"
        document_id = run.json()["answer"].rsplit("/", 2)[-2]

        unauthorized = await client.get(f"/v1/documents/{document_id}/download")
        assert unauthorized.status_code == 401
        download = await client.get(f"/v1/documents/{document_id}/download", headers=headers)

    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    doc = Document(io.BytesIO(download.content))
    assert [paragraph.text for paragraph in doc.paragraphs] == [
        "Compte rendu", "Décision principale", "Prochaine étape"
    ]


@pytest.mark.asyncio
async def test_document_edit_requires_approval_then_replaces_content(tmp_path):
    create_brain = ScriptedBrain(RequestedAction(
        name="document_create",
        arguments={"title": "Plan initial", "content": "Ancienne version"},
    ))
    app = create_app(Settings(database_path=tmp_path / "emefa.db"), brain=create_brain)
    token = app.state.devices.enroll("Claude")[1]
    headers = {"Authorization": f"Bearer {token}"}

    async with await _client_for(app) as client:
        created = await client.post("/v1/agent/runs", json={"message": "Crée le plan"}, headers=headers)
        document_id = created.json()["answer"].rsplit("/", 2)[-2]

        edit_brain = ScriptedBrain(RequestedAction(
            name="document_edit",
            arguments={"document_id": document_id, "title": "Plan final", "content": "Nouvelle version"},
        ))
        app.state.agent.brain = edit_brain
        pending = await client.post("/v1/agent/runs", json={"message": "Modifie le plan"}, headers=headers)
        assert pending.json()["status"] == "confirmation_required"
        action_id = pending.json()["action_id"]

        approved = await client.post(
            f"/v1/agent/approvals/{action_id}/decision",
            json={"approve": True}, headers=headers,
        )
        assert approved.json()["status"] == "completed"
        download = await client.get(f"/v1/documents/{document_id}/download", headers=headers)

    doc = Document(io.BytesIO(download.content))
    assert [paragraph.text for paragraph in doc.paragraphs] == ["Plan final", "Nouvelle version"]
