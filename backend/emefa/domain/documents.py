"""Persistent, path-safe Word document artifacts for EMEFA."""

from __future__ import annotations

import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

from docx import Document
from docx.document import Document as DocxDocument

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: Office artefacts EMEFA can produce, and what a browser should call them.
#: The store keeps them side by side because they are the same thing to the
#: user: something EMEFA made, in the Livrables list, with a download link.
ARTIFACT_MIME_TYPES: dict[str, str] = {
    "docx": _DOCX_MIME,
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "pdf": "application/pdf",
    "csv": "text/csv",
}


class DocumentNotFoundError(LookupError):
    pass


class DocumentStore:
    """Store DOCX artifacts under one server-controlled persistent directory."""

    mime_type = _DOCX_MIME

    def __init__(self, database_path: Path) -> None:
        self.root = Path(database_path).parent / "documents"
        self.root.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _validated_id(document_id: str) -> str:
        try:
            value = str(UUID(str(document_id)))
        except (TypeError, ValueError, AttributeError) as exc:
            raise DocumentNotFoundError("document_not_found") from exc
        return value

    def _path(self, document_id: str, extension: str = "docx") -> Path:
        return self.root / f"{self._validated_id(document_id)}.{extension}"

    def _sidecar(self, document_id: str) -> Path:
        return self.root / f"{self._validated_id(document_id)}.json"

    def _locate(self, document_id: str) -> Path:
        """Find the artefact whatever format it is.

        Word documents predate the other formats and carry no sidecar, so the
        `.docx` path is tried first and the rest are found by scanning. That
        keeps every document created before spreadsheets existed readable.
        """
        identifier = self._validated_id(document_id)
        direct = self._path(identifier)
        if direct.is_file():
            return direct
        for extension in ARTIFACT_MIME_TYPES:
            candidate = self._path(identifier, extension)
            if candidate.is_file():
                return candidate
        raise DocumentNotFoundError("document_not_found")

    @staticmethod
    def _clean_title(title: object) -> str:
        value = str(title).strip()[:180]
        return value or "Document EMEFA"

    @staticmethod
    def _clean_content(content: object) -> str:
        return str(content).strip()[:100_000]

    @staticmethod
    def _build(title: str, content: str) -> DocxDocument:
        document = Document()
        document.add_heading(title, level=0)
        for line in content.splitlines():
            document.add_paragraph(line)
        return document

    def _save(self, document: DocxDocument, destination: Path) -> None:
        temporary = destination.with_suffix(".tmp.docx")
        document.save(str(temporary))
        os.replace(temporary, destination)

    def create(self, title: object, content: object) -> dict[str, str]:
        document_id = str(uuid4())
        clean_title = self._clean_title(title)
        destination = self._path(document_id)
        self._save(self._build(clean_title, self._clean_content(content)), destination)
        return self.describe(document_id)

    def edit(self, document_id: str, title: object | None, content: object) -> dict[str, str]:
        destination = self._path(document_id)
        if not destination.is_file():
            raise DocumentNotFoundError("document_not_found")
        existing = Document(str(destination))
        current_title = existing.paragraphs[0].text if existing.paragraphs else "Document EMEFA"
        clean_title = self._clean_title(current_title if title is None else title)
        self._save(self._build(clean_title, self._clean_content(content)), destination)
        return self.describe(document_id)

    def save_artifact(
        self, data: bytes, extension: str, title: object, kind: str = ""
    ) -> dict[str, Any]:
        """Store a workbook, a deck or any other office artefact.

        The title is written to a sidecar rather than read back out of the
        file: a spreadsheet has no first paragraph to read, and re-opening a
        deck to learn its own name would be absurd.
        """
        extension = extension.strip().lower().lstrip(".")
        if extension not in ARTIFACT_MIME_TYPES:
            raise ValueError(f"unsupported artifact format: {extension!r}")
        document_id = str(uuid4())
        destination = self._path(document_id, extension)
        temporary = destination.with_suffix(f".tmp.{extension}")
        temporary.write_bytes(data)
        os.replace(temporary, destination)
        self._sidecar(document_id).write_text(
            json.dumps(
                {
                    "title": self._clean_title(title),
                    "extension": extension,
                    "kind": kind,
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return self.describe(document_id)

    def get(self, document_id: str) -> Path:
        return self._locate(document_id)

    def describe(self, document_id: str) -> dict[str, Any]:
        path = self._locate(document_id)
        extension = path.suffix.lstrip(".")
        kind = ""
        sidecar = self._sidecar(document_id)
        if sidecar.is_file():
            try:
                metadata = json.loads(sidecar.read_text(encoding="utf-8"))
                title = str(metadata.get("title") or "Document EMEFA")
                kind = str(metadata.get("kind") or "")
            except (OSError, json.JSONDecodeError):
                title = "Document EMEFA"
        elif extension == "docx":
            document = Document(str(path))
            title = document.paragraphs[0].text if document.paragraphs else "Document EMEFA"
        else:
            title = "Document EMEFA"
        slug = re.sub(r"[^A-Za-z0-9À-ÿ]+", "-", title).strip("-")[:80] or "document-emefa"
        modified_at = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc).isoformat()
        return {
            "document_id": self._validated_id(document_id),
            "title": title,
            "kind": kind,
            "filename": f"{slug}.{extension}",
            "content_type": ARTIFACT_MIME_TYPES.get(extension, self.mime_type),
            "size_bytes": path.stat().st_size,
            "updated_at": modified_at,
            "download_url": f"/v1/documents/{self._validated_id(document_id)}/download",
        }

    def list(self, limit: int = 100) -> list[dict[str, Any]]:
        """Return newest generated artifacts without trusting filenames as IDs."""

        candidates = sorted(
            (
                path
                for extension in ARTIFACT_MIME_TYPES
                for path in self.root.glob(f"*.{extension}")
                if ".tmp." not in path.name
            ),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        records: list[dict[str, Any]] = []
        for path in candidates[: max(1, min(limit, 500))]:
            try:
                records.append(self.describe(path.stem))
            except (DocumentNotFoundError, ValueError, OSError):
                continue
        return records
